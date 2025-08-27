import json
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Any, Dict

def set_run_font(run, font_name, font_size, bold=False, color=None):
    """Apply font style and color to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)
    # Set line height and space before for the paragraph containing this run
    para = run._element.getparent()
    if hasattr(run, 'paragraph'):
        para = run.paragraph
    if para is not None and hasattr(para, 'paragraph_format'):
        para.paragraph_format.line_spacing = 1.1  # Smaller line height
        para.paragraph_format.space_before = Pt(3)

def set_table_header_bg_color(cell, rgb_hex="EDEDED"):
    """Set the background color of a table header cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), rgb_hex)
    tc_pr.append(shd)

def format_label(label):
    """Format JSON keys into human-readable labels."""
    return label.replace("_", " ").capitalize()

# -------------------------------
# Table Creation
# -------------------------------

def add_table(doc, title, data, columns):
    heading = doc.add_heading(level=3)
    run = heading.add_run(title)
    set_run_font(run, "Cambria", 15, bold=True, color=(26, 35, 126))
    heading.paragraph_format.space_after = Pt(12)

    if not data or not isinstance(data, list):
        p = doc.add_paragraph("No data available.")
        set_run_font(p.runs[0], "Cambria", 11)
        return

    formatted_columns = [format_label(col) for col in columns]
    table = doc.add_table(rows=1, cols=len(formatted_columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(formatted_columns):
        cell_run = hdr_cells[i].paragraphs[0].add_run(col)
        set_run_font(cell_run, "Cambria", 10, bold=True, color=(255,255,255))
        set_table_header_bg_color(hdr_cells[i], "17365c")
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].paragraph_format.line_spacing = 1.5
        hdr_cells[i].paragraphs[0].paragraph_format.space_before = Pt(6)

    # Fill table rows
    for row in data:
        if not isinstance(row, dict):
            continue
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            value = row.get(col, "")
            # If value is False, empty list, or None, use empty string
            if value is False or value == [] or value is None:
                value = ""
            if isinstance(value, float):
                value = f"{value:.2f}"
            # Special style for total prices
            if str(col).lower() in {"total", "total_amount", "totalcost", "total price", "total price (€)", "total_cost", "totalcosts", "application_price"}:
                run = row_cells[i].paragraphs[0].add_run(str(value))
                set_run_font(run, "Arial", 10, bold=True)
            else:
                run = row_cells[i].paragraphs[0].add_run(str(value))
                set_run_font(run, "Cambria", 11)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row_cells[i].paragraphs[0].paragraph_format.line_spacing = 1.5
            row_cells[i].paragraphs[0].paragraph_format.space_before = Pt(6)

# -------------------------------
# Cost Summary Section
# -------------------------------

def add_cost_summary(doc, cost_summary):
    heading = doc.add_heading("Internal Cost Summary", level=2)
    set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
    heading.paragraph_format.space_after = Pt(12)

    for key, value in cost_summary.items():
        label = format_label(key)
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        set_run_font(run_label, "Cambria", 13, bold=True)
        run_value = p.add_run(f"{value}")
        set_run_font(run_value, "Cambria", 13, color=(26, 35, 126))  # Highlight totals in dark blue

# -------------------------------
# Main Report Generation
# -------------------------------

def generate_internal_costs_doc(data: Dict[str, Any], output_path: str = "Internal_Costs_Report.docx"):
    doc = Document()
    # Set margins
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    internal_costs = data.get("internalCosts", {})

    # Title Page
    heading = doc.add_heading(level=1)
    run = heading.add_run(internal_costs.get("offer_title", "Internal Company Costs"))
    set_run_font(run, "Cambria", 15, bold=True, color=(26, 35, 126))
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(18)
    doc.add_paragraph("")

    # Introduction and context
    intro = internal_costs.get("cost_description", "Official internal cost report for the company, based on the construction proposal. All values are in EUR unless otherwise specified.")
    p = doc.add_paragraph(intro)
    set_run_font(p.runs[0], "Cambria", 12)
    doc.add_paragraph("")


    # --- Area Summaries ---
    site_area_summary = internal_costs.get("site_area_summary", [])
    for area in site_area_summary:
        heading = doc.add_heading(f"Area: {area.get('area', '')}", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        # Area summary as list
        area_summary_cols = ["total_cost", "markup_percentage", "final_cost_for_client_eur", "materials", "labor", "subcontractors", "equipment"]
        area_summary_labels = ["Totale", "Markup (%)", "Totale Cliente", "Materiali", "Manodopera", "Subappalti", "Attrezzature"]
        p = doc.add_paragraph()
        for k, label in zip(area_summary_cols, area_summary_labels):
            value = area.get(k, "")
            if value:
                run = p.add_run(f"{label}: {value}, ")
                set_run_font(run, "Cambria", 11)
        # Resource types
        resource_types = area.get("resource_types", [])
        if resource_types:
            p = doc.add_paragraph("Tipi di risorse: " + ", ".join(resource_types))
            set_run_font(p.runs[0], "Cambria", 11)
        # Resources as list
        resources = area.get("resources", [])
        if resources:
            p = doc.add_paragraph("Risorse:")
            set_run_font(p.runs[0], "Cambria", 11, bold=True)
            for r in resources:
                q = r.get("quantity", "")
                unit = r.get("unit", "")
                unit_price = r.get("unitPrice", "")
                total = r.get("totalPrice", "")
                name = r.get("name", "")
                typ = r.get("type", "")
                line = f"- {name} ({typ}), {q} {unit}, {unit_price} / {unit}, Totale: {total}"
                p = doc.add_paragraph(line)
                set_run_font(p.runs[0], "Cambria", 10)
        # Work activities as list
        work_activities = area.get("work_activities", [])
        if work_activities:
            p = doc.add_paragraph("Attività di lavoro:")
            set_run_font(p.runs[0], "Cambria", 11, bold=True)
            for w in work_activities:
                desc = w.get("description", "")
                q = w.get("quantity", "")
                unit = w.get("unit", "")
                unit_price = w.get("unitPrice", "")
                total = w.get("totalPrice", "")
                line = f"- {desc}, {q} {unit}, {unit_price} / {unit}, Totale: {total}"
                p = doc.add_paragraph(line)
                set_run_font(p.runs[0], "Cambria", 10)
        doc.add_paragraph("")

    # --- Offerta economica e dettaglio costi per aree ---
    # (Title and intro already added at the top, do not repeat here)

    aree = internal_costs.get("areas", [])
    for area in aree:
        nome_area = area.get("name", "Area")
        totale = area.get("total", "")
        markup = area.get("markup", "")
        totale_cliente = area.get("total_client", "")
        materiali = area.get("materials", "")
        manodopera = area.get("labor", "")
        subappalti = area.get("subcontracts", "")
        attrezzature = area.get("equipment", "")
        risorse = area.get("resources", [])
        attivita = area.get("work_activities", [])

        doc.add_heading(f"Area: {nome_area}", level=2)
        p = doc.add_paragraph()
        p.add_run(f"Totale: {totale}").bold = True
        p.add_run(f", Markup (%): {markup}")
        p.add_run(f", Totale Cliente: {totale_cliente}")
        p.add_run(f", Materiali: {materiali}")
        p.add_run(f", Manodopera: {manodopera}")
        p.add_run(f", Subappalti: {subappalti}")
        p.add_run(f", Attrezzature: {attrezzature}")

        # Tipi di risorse
        tipi_risorse = set(r.get("type_label", r.get("type", "")) for r in risorse)
        if tipi_risorse:
            doc.add_paragraph("Tipi di risorse: " + ", ".join(sorted(tipi_risorse)))

        # Tabella risorse
        if risorse:
            doc.add_paragraph("Risorse:")
            risorse_cols = ["nome", "tipo", "quantita", "unita", "prezzo_unitario", "totale"]
            risorse_labels = ["Nome", "Tipo", "Quantità", "Unità", "Prezzo unitario", "Totale"]
            table = doc.add_table(rows=1, cols=len(risorse_cols))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, label in enumerate(risorse_labels):
                run = hdr_cells[i].paragraphs[0].add_run(label)
                set_run_font(run, "Cambria", 10, bold=True, color=(255,255,255))
                set_table_header_bg_color(hdr_cells[i], "17365c")
            for r in risorse:
                row_cells = table.add_row().cells
                row_cells[0].text = r.get("name", "")
                row_cells[1].text = r.get("type_label", r.get("type", ""))
                row_cells[2].text = str(r.get("quantity", ""))
                row_cells[3].text = r.get("unit", "")
                row_cells[4].text = str(r.get("unit_price", ""))
                row_cells[5].text = str(r.get("total", ""))
                for cell in row_cells:
                    cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                    cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(8)
                    for run in cell.paragraphs[0].runs:
                        set_run_font(run, "Cambria", 9)
            doc.add_paragraph("")

        # Attività di lavoro
        if attivita:
            doc.add_paragraph("Attività di lavoro:")
            attivita_cols = ["descrizione", "quantita", "unita", "prezzo_unitario", "totale"]
            attivita_labels = ["Descrizione", "Quantità", "Unità", "Prezzo unitario", "Totale"]
            table = doc.add_table(rows=1, cols=len(attivita_cols))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, label in enumerate(attivita_labels):
                run = hdr_cells[i].paragraphs[0].add_run(label)
                set_run_font(run, "Cambria", 10, bold=True, color=(255,255,255))
                set_table_header_bg_color(hdr_cells[i], "17365c")
            for a in attivita:
                row_cells = table.add_row().cells
                row_cells[0].text = a.get("description", "")
                row_cells[1].text = str(a.get("quantity", ""))
                row_cells[2].text = a.get("unit", "")
                row_cells[3].text = str(a.get("unit_price", ""))
                row_cells[4].text = str(a.get("total", ""))
                for cell in row_cells:
                    cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                    cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(8)
                    for run in cell.paragraphs[0].runs:
                        set_run_font(run, "Cambria", 9)
            doc.add_paragraph("")
    materials_list = internal_costs.get("materialsList", [])
    if materials_list:
        heading = doc.add_heading("Lista Materiali", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        mat_cols = ["item", "quantity", "unit", "unitPrice", "total_quantity", "provider_name", "price_of_unity_provider", "total_price", "company_cost_eur", "markup_percentage", "final_cost_for_client_eur"]
        mat_labels = ["Articolo", "Quantità", "Unità", "Prezzo unitario", "Quantità totale", "Fornitore", "Prezzo unità fornitore", "Prezzo totale", "Costo azienda (€)", "Percentuale ricarico", "Prezzo finale cliente (€)"]
        # Group by 5 items per table, always 5 columns (4 items + 1 label), fill empty if needed
        for group_start in range(0, len(materials_list), 5):
            group = materials_list[group_start:group_start+5]
            table = doc.add_table(rows=len(mat_cols), cols=6)
            table.style = "Table Grid"
            for i, (col, label) in enumerate(zip(mat_cols, mat_labels)):
                hdr_cell = table.cell(i, 0)
                hdr_run = hdr_cell.paragraphs[0].add_run(label)
                set_run_font(hdr_run, "Cambria", 10, bold=True, color=(255,255,255))
                set_table_header_bg_color(hdr_cell, "17365c")
                hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                hdr_cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                hdr_cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                hdr_cell.paragraphs[0].paragraph_format.space_after = Pt(8)
            for j in range(5):
                if j < len(group):
                    row = group[j]
                    for i, col in enumerate(mat_cols):
                        cell = table.cell(i, j+1)
                        cell.text = str(row.get(col, ""))
                        cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                        cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                        cell.paragraphs[0].paragraph_format.space_after = Pt(8)
                        for run in cell.paragraphs[0].runs:
                            set_run_font(run, "Cambria", 9)
                else:
                    for i in range(len(mat_cols)):
                        cell = table.cell(i, j+1)
                        cell.text = ""
            doc.add_paragraph("")

    # --- Personale ---
    personnel = internal_costs.get("personnel", [])
    if personnel:
        heading = doc.add_heading("Personale", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        from collections import defaultdict
        # Split personnel by type
        employees = [p for p in personnel if p.get("type", "").lower() == "dipendente"]
        subcontractors = [p for p in personnel if p.get("type", "").lower() == "subappaltatore"]

        def group_by_site_works(personnel_list):
            works_map = defaultdict(list)
            for p in personnel_list:
                sw = p.get('site_works', [])
                if isinstance(sw, list):
                    for w in sw:
                        if isinstance(w, dict):
                            key = f"{w.get('type', '')} ({w.get('category', '')})"
                        else:
                            key = str(w)
                        works_map[key].append(p)
                else:
                    works_map[str(sw)].append(p)
            return works_map

        # Employees table(s)
        if employees:
            heading_emp = doc.add_heading("Dipendenti", level=3)
            set_run_font(heading_emp.runs[0], "Cambria", 13, bold=True, color=(26, 35, 126))
            emp_cols = ["role", "count", "duration", "price_per_hour", "total"]
            emp_labels = ["Ruolo", "Numero", "Durata", "Prezzo orario", "Totale"]
            emp_works_map = group_by_site_works(employees)
            for work, group in emp_works_map.items():
                subheading = doc.add_heading(f"Lavoro: {work}", level=4)
                set_run_font(subheading.runs[0], "Cambria", 12, bold=True, color=(26, 35, 126))
                for group_start in range(0, len(group), 5):
                    subgroup = group[group_start:group_start+5]
                    # Always 5 columns for items, fill empty if needed
                    table = doc.add_table(rows=len(emp_cols), cols=6)
                    table.style = "Table Grid"
                    for i, (col, label) in enumerate(zip(emp_cols, emp_labels)):
                        hdr_cell = table.cell(i, 0)
                        hdr_run = hdr_cell.paragraphs[0].add_run(label)
                        set_run_font(hdr_run, "Cambria", 10, bold=True, color=(255,255,255))
                        set_table_header_bg_color(hdr_cell, "17365c")
                        hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        hdr_cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                        hdr_cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                        hdr_cell.paragraphs[0].paragraph_format.space_after = Pt(8)
                    for j in range(5):
                        if j < len(subgroup):
                            p = subgroup[j]
                            for i, col in enumerate(emp_cols):
                                val = p.get(col, "")
                                cell = table.cell(i, j+1)
                                cell.text = str(val)
                                cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                                cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                                cell.paragraphs[0].paragraph_format.space_after = Pt(8)
                                for run in cell.paragraphs[0].runs:
                                    set_run_font(run, "Cambria", 9)
                        else:
                            for i in range(len(emp_cols)):
                                cell = table.cell(i, j+1)
                                cell.text = ""
                    doc.add_paragraph("")

        # Subcontractors table(s)
        if subcontractors:
            heading_sub = doc.add_heading("Subappaltatori", level=3)
            set_run_font(heading_sub.runs[0], "Cambria", 13, bold=True, color=(26, 35, 126))
            sub_cols = ["role", "unit_measure", "price_per_unit", "quantity", "total"]
            sub_labels = ["Ruolo", "Unità di misura", "Prezzo unitario", "Quantità", "Totale"]
            sub_works_map = group_by_site_works(subcontractors)
            for work, group in sub_works_map.items():
                subheading = doc.add_heading(f"Lavoro: {work}", level=4)
                set_run_font(subheading.runs[0], "Cambria", 12, bold=True, color=(26, 35, 126))
                for group_start in range(0, len(group), 5):
                    subgroup = group[group_start:group_start+5]
                    # Always 5 columns for items, fill empty if needed
                    table = doc.add_table(rows=len(sub_cols), cols=6)
                    table.style = "Table Grid"
                    for i, (col, label) in enumerate(zip(sub_cols, sub_labels)):
                        hdr_cell = table.cell(i, 0)
                        hdr_run = hdr_cell.paragraphs[0].add_run(label)
                        set_run_font(hdr_run, "Cambria", 10, bold=True, color=(255,255,255))
                        set_table_header_bg_color(hdr_cell, "17365c")
                        hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        hdr_cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                        hdr_cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                        hdr_cell.paragraphs[0].paragraph_format.space_after = Pt(8)
                    for j in range(5):
                        if j < len(subgroup):
                            p = subgroup[j]
                            for i, col in enumerate(sub_cols):
                                val = p.get(col, "")
                                cell = table.cell(i, j+1)
                                cell.text = str(val)
                                cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                                cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                                cell.paragraphs[0].paragraph_format.space_after = Pt(8)
                                for run in cell.paragraphs[0].runs:
                                    set_run_font(run, "Cambria", 9)
                        else:
                            for i in range(len(sub_cols)):
                                cell = table.cell(i, j+1)
                                cell.text = ""
                    doc.add_paragraph("")

        # Safety courses requirements table
        all_safety = set()
        for p in personnel:
            sc = p.get('safety_courses_requirements', [])
            if isinstance(sc, list):
                all_safety.update(sc)
            elif sc:
                all_safety.add(str(sc))
        if all_safety:
            heading = doc.add_heading("Requisiti corsi di sicurezza", level=3)
            set_run_font(heading.runs[0], "Cambria", 13, bold=True, color=(26, 35, 126))
            table = doc.add_table(rows=len(all_safety), cols=1)
            table.style = "Table Grid"
            for i, req in enumerate(sorted(all_safety)):
                cell = table.cell(i, 0)
                cell.text = req
                cell.paragraphs[0].paragraph_format.line_spacing = 1.3
                cell.paragraphs[0].paragraph_format.space_before = Pt(8)
                cell.paragraphs[0].paragraph_format.space_after = Pt(8)
                for run in cell.paragraphs[0].runs:
                    set_run_font(run, "Cambria", 9)
            doc.add_paragraph("")

    # --- Logistica ---
    logistics = internal_costs.get("logistics", [])
    if logistics:
        heading = doc.add_heading("Logistica", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        log_data = []
        for l in logistics:
            row = dict(l)
            if isinstance(row.get("site_category"), list):
                row["site_category"] = ", ".join(row["site_category"])
            if isinstance(row.get("site_works"), list):
                row["site_works"] = ", ".join([f"{w.get('type', w)} ({w.get('category', '')})" if isinstance(w, dict) else str(w) for w in row["site_works"]])
            log_data.append(row)
        log_cols = ["description", "duration", "unity", "total_price", "site_category", "site_works"]
        add_table(doc, "", log_data, log_cols)

    # --- Costi Diretti ---
    direct_costs = internal_costs.get("direct_costs", [])
    if direct_costs:
        heading = doc.add_heading("Costi Diretti", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        direct_cols = ["category", "description", "unit", "price", "total_price"]
        add_table(doc, "", direct_costs, direct_cols)

    # --- Costi Indiretti ---
    indirect_costs = internal_costs.get("indirect_costs", [])
    if indirect_costs:
        heading = doc.add_heading("Costi Indiretti", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        indirect_cols = ["category", "description", "unit", "price", "total_price"]
        add_table(doc, "", indirect_costs, indirect_cols)

    # --- Attrezzature ---
    equipment = internal_costs.get("equipment", [])
    if equipment:
        heading = doc.add_heading("Attrezzature", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        eq_cols = ["name", "quantity", "unity", "price_per_unit", "price_of_unity_provider", "company_cost_eur", "markup_percentage", "final_cost_for_client_eur"]
        add_table(doc, "", equipment, eq_cols)

    # --- Price Summary ---
    price_summary = internal_costs.get("price_summary", {})
    if price_summary:
        heading = doc.add_heading("Price Summary", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        # Main values as list, bold total labels and prices, add currency
        summary_keys = ["subtotal_price", "global_costs", "company_profit", "margin_check", "markup", "rounding", "total_costs", "application_price"]
        summary_labels = ["Subtotale", "Costi Globali", "Profitto Azienda", "Controllo Margine", "Markup", "Arrotondamento", "Costi Totali", "Prezzo Applicato"]
        currency = internal_costs.get("currency", "EUR").upper()
        for k, label in zip(summary_keys, summary_labels):
            value = price_summary.get(k, "")
            if value:
                p = doc.add_paragraph()
                run_label = p.add_run(f"{label}: ")
                run_price = p.add_run(f"{value} {currency}")
                # Bold for totals and prices
                if k in ["subtotal_price", "total_costs", "application_price"]:
                    run_label.bold = True
                    run_price.bold = True
                set_run_font(run_label, "Cambria", 11)
                set_run_font(run_price, "Cambria", 11)
        # Explanations
        explanations = price_summary.get("explanation_of_summary", {})
        if explanations:
            p = doc.add_paragraph("Spiegazione dei campi del riepilogo:")
            set_run_font(p.runs[0], "Cambria", 11, bold=True)
            for k, v in explanations.items():
                label = k.replace("_", " ").capitalize()
                p = doc.add_paragraph(f"- {label}: {v}")
                set_run_font(p.runs[0], "Cambria", 10)
        # Summary by category as list
        summary_by_cat = price_summary.get("summary_by_category", {})
        if summary_by_cat:
            p = doc.add_paragraph("Riepilogo per categoria:")
            set_run_font(p.runs[0], "Cambria", 11, bold=True)
            for k, v in summary_by_cat.items():
                label = k.replace("_", " ").capitalize()
                line = f"- {label}: {v}"
                para = doc.add_paragraph(line)
                set_run_font(para.runs[0], "Cambria", 10)

    # --- Depreciation ---
    depreciation = internal_costs.get("deprecation_fixed_amount", [])
    if depreciation:
        heading = doc.add_heading("Depreciation", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        for d in depreciation:
            comp = d.get("component", "")
            dep = d.get("depreciation_eur", "")
            perc = d.get("percentage", "")
            rounding = d.get("rounding", "")
            line = f"- {comp}, Ammortamento: {dep}, Percentuale: {perc}, Arrotondamento: {rounding}"
            para = doc.add_paragraph(line)
            set_run_font(para.runs[0], "Cambria", 10)

    # --- Risk Analysis ---
    risk_analysis = internal_costs.get("risk_analysis", {})
    if risk_analysis:
        heading = doc.add_heading("Risk Analysis", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        for k, v in risk_analysis.items():
            p = doc.add_paragraph(f"{format_label(k)}: {v}")
            set_run_font(p.runs[0], "Cambria", 11)

    # --- Cronoprogramma e Gantt ---
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
    import datetime
    import tempfile
    schedule = internal_costs.get("projectSchedule", [])
    if schedule:
        heading = doc.add_heading("Cronoprogramma dei Lavori", level=2)
        set_run_font(heading.runs[0], "Cambria", 15, bold=True, color=(26, 35, 126))
        heading.paragraph_format.space_after = Pt(12)
        # Gantt chart with matplotlib
        # Prepare data
        activities = []
        starts = []
        ends = []
        for entry in schedule:
            activities.append(entry.get("activity", ""))
            try:
                starts.append(int(entry.get("starting", 0)))
                ends.append(int(entry.get("finishing", 0)))
            except Exception:
                starts.append(0)
                ends.append(0)
        # Use days as offsets from a base date
        base_date = datetime.date.today()
        start_dates = [base_date + datetime.timedelta(days=s-1) for s in starts]
        end_dates = [base_date + datetime.timedelta(days=e-1) for e in ends]
        fig, ax = plt.subplots(figsize=(6, 0.3*len(activities)+1), dpi=80)
        for i, (act, start, end) in enumerate(zip(activities, start_dates, end_dates)):
            ax.barh(act, (end-start).days+1, left=start, height=0.5, color="#4F81BD")
        # Remove x-axis labels
        ax.set_xlabel("")
        ax.set_xticklabels([])
        plt.ylabel('Attività')
        plt.tight_layout()
        # Save to temp file (after chart is drawn)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
            plt.savefig(tmpfile.name, bbox_inches='tight', dpi=80)
            plt.close(fig)
            doc.add_picture(tmpfile.name, width=Inches(6))
        # Add a readable list of activity type and days required below the Gantt chart
        doc.add_paragraph("Attività e giorni richiesti:")
        for entry in schedule:
            activity = entry.get("activity", "")
            start = entry.get("starting", 0)
            end = entry.get("finishing", 0)
            try:
                days = int(end) - int(start) + 1
            except Exception:
                days = "-"
            doc.add_paragraph(f"- {activity}: {days} giorni", style="List Bullet")

    # Save file
    doc.save(output_path)
    return output_path
